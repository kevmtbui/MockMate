import os
from typing import Optional, Dict
import PyPDF2
from docx import Document
import io
import json
import google.generativeai as genai
from config.prompts import AIPrompts

class ResumeService:
    @staticmethod
    def extract_text_from_file(file_path: str) -> Optional[str]:
        """Extract text from PDF or DOCX file"""
        try:
            if file_path.lower().endswith('.pdf'):
                return ResumeService._extract_from_pdf(file_path)
            elif file_path.lower().endswith(('.doc', '.docx')):
                return ResumeService._extract_from_docx(file_path)
            else:
                return None
        except Exception as e:
            print(f"Error extracting text from file: {e}")
            return None
    
    @staticmethod
    def extract_text_from_upload(file_content: bytes, filename: str) -> Optional[str]:
        """Extract text from uploaded file content"""
        try:
            if filename.lower().endswith('.pdf'):
                return ResumeService._extract_from_pdf_content(file_content)
            elif filename.lower().endswith(('.doc', '.docx')):
                return ResumeService._extract_from_docx_content(file_content)
            else:
                return None
        except Exception as e:
            print(f"Error extracting text from upload: {e}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    @staticmethod
    def _extract_from_pdf_content(file_content: bytes) -> str:
        """Extract text from PDF content"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    @staticmethod
    def _extract_from_docx_content(file_content: bytes) -> str:
        """Extract text from DOCX content"""
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    @staticmethod
    def summarize_resume(resume_text: str) -> str:
        """Create a summary of the resume for AI context"""
        # This could be enhanced with AI summarization
        # For now, return a truncated version
        if len(resume_text) > 1000:
            return resume_text[:1000] + "..."
        return resume_text
    
    @staticmethod
    def analyze_resume_with_ai(resume_text: str) -> Dict:
        """Use Gemini AI to analyze and extract key information from resume"""
        try:
            # Configure Gemini API
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                print("GEMINI_API_KEY not found, using fallback analysis")
                return ResumeService._get_fallback_analysis(resume_text)
            
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            prompt = AIPrompts.RESUME_ANALYSIS.format(resume_text=resume_text)
            
            print(f"Calling Gemini for resume analysis with prompt length: {len(prompt)}")
            response = model.generate_content(prompt)
            response_text = response.text
            
            print(f"Gemini resume analysis response received, length: {len(response_text)}")
            
            # Try to extract JSON from response
            try:
                # Find JSON object in response
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx != -1:
                    json_text = response_text[start_idx:end_idx]
                    analysis = json.loads(json_text)
                    print(f"Resume analysis completed successfully. Keys: {list(analysis.keys())}")
                    return analysis
                else:
                    raise ValueError("No JSON object found in response")
                    
            except (json.JSONDecodeError, ValueError) as e:
                print(f"JSON parsing error in resume analysis: {e}")
                print(f"Raw response: {response_text}")
                return ResumeService._get_fallback_analysis(resume_text)
                
        except Exception as e:
            print(f"Error in Gemini resume analysis: {e}")
            return ResumeService._get_fallback_analysis(resume_text)
    
    @staticmethod
    def _get_fallback_analysis(resume_text: str) -> Dict:
        """Fallback analysis if AI fails"""
        return {
            "name": "Candidate",
            "email": None,
            "phone": None,
            "location": None,
            "summary": "Professional with relevant experience in the field.",
            "experience_years": 3,
            "current_role": "Professional",
            "current_company": "Company",
            "education": "Degree in relevant field",
            "skills": ["Problem Solving", "Communication", "Teamwork"],
            "technologies": ["General Technology"],
            "industries": ["Technology"],
            "achievements": ["Professional accomplishments"],
            "certifications": [],
            "languages": ["English"],
            "experience_level": "mid",
            "key_strengths": ["Adaptability", "Learning", "Collaboration"],
            "career_progression": "Steady professional growth"
        }
    
    @staticmethod
    def generate_resume_summary_for_ai(resume_text: str, analysis: Dict = None) -> str:
        """Generate a comprehensive summary for AI question generation"""
        if not analysis:
            analysis = ResumeService.analyze_resume_with_ai(resume_text)
        
        summary = f"""
        CANDIDATE PROFILE:
        Name: {analysis.get('name', 'Candidate')}
        Experience Level: {analysis.get('experience_level', 'mid')} level
        Total Experience: {analysis.get('experience_years', 3)} years
        Current Role: {analysis.get('current_role', 'Professional')} at {analysis.get('current_company', 'Company')}
        Education: {analysis.get('education', 'Relevant degree')}
        
        KEY SKILLS: {', '.join(analysis.get('skills', []))}
        TECHNOLOGIES: {', '.join(analysis.get('technologies', []))}
        INDUSTRIES: {', '.join(analysis.get('industries', []))}
        
        PROFESSIONAL SUMMARY: {analysis.get('summary', 'Experienced professional')}
        
        KEY ACHIEVEMENTS: {', '.join(analysis.get('achievements', []))}
        STRENGTHS: {', '.join(analysis.get('key_strengths', []))}
        
        CAREER PROGRESSION: {analysis.get('career_progression', 'Steady professional growth')}
        """
        
        return summary.strip()
